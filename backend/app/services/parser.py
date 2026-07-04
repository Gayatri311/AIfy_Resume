import re
from pathlib import Path
from typing import Optional
from docx import Document
from app.schemas.resume import ResumeData, PersonalDetails, ExperienceItem, EducationItem, ProjectItem
from app.services.bullet_utils import (
    is_bullet_line,
    normalize_bullet_list,
    split_inline_bullets,
    split_merged_accomplishments,
    strip_bullet_prefix,
)
from app.services.resume_format import normalize_resume_data
from app.services.text_cleanup import preprocess_resume_text, split_name_contact_line
from app.services.section_patterns import match_section_line
from app.services.pdf_extractor import extract_text_from_pdf

SECTION_HEADERS = {
    "summary": re.compile(
        r"^(professional\s+)?(summary|profile|objective|about(\s+me)?)$", re.I
    ),
    "experience": re.compile(
        r"^(professional\s+)?(work\s+)?experience(\s+history)?$|^(employment(\s+history)?|work\s+history)$",
        re.I,
    ),
    "projects": re.compile(r"^(personal\s+)?projects(\s+|$)|^key\s+projects$", re.I),
    "skills": re.compile(
        r"^(technical\s+)?skills$|^core\s+competencies$|^key\s+competencies$|^technologies$",
        re.I,
    ),
    "education": re.compile(
        r"^education(?:\s*&\s*certifications?)?$|^academic(\s+background)?$", re.I
    ),
    "certifications": re.compile(r"^certifications?$|^licenses?$", re.I),
    "awards": re.compile(
        r"^(key\s+achievements(?:\s+and\s+interests)?|achievements?|awards?|honors?)$", re.I
    ),
    "languages": re.compile(r"^languages$", re.I),
    "interests": re.compile(r"^interests$", re.I),
}

DATE_RANGE = re.compile(
    r"(\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\b|\b\d{1,2}/\d{4}\b|\b\d{4}\b)"
    r"\s*[-–—to]+\s*"
    r"(\b(?:present|current|now)\b|\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\b|\b\d{1,2}/\d{4}\b|\b\d{4}\b)",
    re.I,
)
BULLET_PREFIX = re.compile(r"^[\u2022\-\*•●▪]\s*")
EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE = re.compile(r"\+?[\d\s().\-]{10,}")
LINKEDIN = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+", re.I)
GITHUB = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-/]+", re.I)
URL_PATTERN = re.compile(r"(?:https?://|www\.)[^\s\)\]>]+", re.I)


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            parts.append("")
            continue
        parts.append(text.rstrip())
    for table in doc.tables:
        parts.append("")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: {ext}")


def _match_section(line: str) -> Optional[str]:
    clean = line.strip().rstrip(":")
    matched = match_section_line(clean)
    if matched:
        return matched
    for name, pattern in SECTION_HEADERS.items():
        if pattern.match(clean):
            return name
    return None


def _is_bullet(line: str) -> bool:
    return is_bullet_line(line)


def _strip_bullet(line: str) -> str:
    return strip_bullet_prefix(line)


def _parse_contact(text: str, personal: PersonalDetails) -> PersonalDetails:
    email = EMAIL.search(text)
    phone = PHONE.search(text)
    linkedin = LINKEDIN.search(text)
    github = GITHUB.search(text)
    if email and not personal.email:
        personal.email = email.group(0)
    if phone and not personal.phone:
        personal.phone = phone.group(0).strip()
    if linkedin and not personal.linkedin:
        personal.linkedin = linkedin.group(0)
    if github and not personal.github:
        personal.github = github.group(0)
    urls = URL_PATTERN.findall(text)
    for url in urls:
        lower = url.lower()
        if "linkedin.com" in lower or "github.com" in lower:
            continue
        if not personal.website:
            personal.website = url
            break
    return personal


def _is_likely_project_title(line: str) -> bool:
    if _is_bullet(line) or line.lower().startswith("tech:"):
        return False
    if URL_PATTERN.search(line) and len(line) < 120:
        return False
    return len(line) < 90 and not line.endswith(".")


def _looks_like_job_header(line: str) -> bool:
    if _is_bullet(line):
        return False
    if DATE_RANGE.search(line):
        return True
    if "|" in line and len(line) < 120:
        return True
    if re.search(r"\b(19|20)\d{2}\b", line) and len(line) < 100:
        return True
    return False


def _parse_education_line(line: str) -> EducationItem:
    line = line.strip()
    if line.upper().startswith("LANGUAGES"):
        line = line.split(None, 1)[1] if len(line.split()) > 1 else line

    year_match = re.search(r"\b(19|20)\d{2}\s*[-–—to]+\s*(19|20)\d{2}\b|\b(19|20)\d{2}\b", line)
    year = ""
    if year_match:
        year = year_match.group(0).replace(" ", "")
        line = line.replace(year_match.group(0), "").strip(" ,-|")

    degree_match = re.search(
        r"\b(M\.?Tech|M\.?S\.?|B\.?Tech|B\.?S\.?|B\.?A\.?|M\.?A\.?|Ph\.?D|MBA|"
        r"Bachelor|Master|Associate)[^,|]*",
        line,
        re.I,
    )
    degree = degree_match.group(0).strip(" ,") if degree_match else ""
    if degree_match:
        line = line.replace(degree_match.group(0), "").strip(" ,-|")

    institution = re.sub(r"\s{2,}", " ", line).strip(" ,-|")
    if not institution and degree:
        institution, degree = degree, ""

    return EducationItem(
        institution=institution or degree,
        degree=degree if institution else "",
        year=year,
    )


def _parse_job_header(line: str) -> tuple[str, str, str, str]:
    line = re.sub(r"\s{2,}", " ", line.strip())

    # Title, MM/YYYY - Present Company Location  (e.g. AI Engineer, 07/2025 - Present Nurix AI - Bangalore)
    title_comma = re.match(
        r"^(.+?),\s*"
        r"(\b(?:\d{1,2}/\d{4}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}|\d{4})\b)\s*"
        r"[-–—to]+\s*"
        r"(Present|Current|Now|\b(?:\d{1,2}/\d{4}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}|\d{4})\b)\s+"
        r"(.+)$",
        line,
        re.I,
    )
    if title_comma:
        return (
            title_comma.group(4).strip(),
            title_comma.group(1).strip(),
            title_comma.group(2).strip(),
            title_comma.group(3).strip(),
        )

    # Role DATE - END Company  (e.g. Product May 2025 – Present MyFi by TIFIN...)
    role_dates_company = re.match(
        r"^(.+?)\s+"
        r"(\b(?:\d{1,2}/\d{4}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4})\b)\s*"
        r"[-–—to]+\s*"
        r"(Present|Current|Now|\b(?:\d{1,2}/\d{4}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4})\b)\s+"
        r"(.+)$",
        line,
        re.I,
    )
    if role_dates_company:
        return (
            role_dates_company.group(4).strip(),
            role_dates_company.group(1).strip(),
            role_dates_company.group(2).strip(),
            role_dates_company.group(3).strip(),
        )

    # Role MM/YYYY - Present  (company on next line)
    role_dates = re.match(
        r"^(.+?)\s+"
        r"(\b(?:\d{1,2}/\d{4}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4})\b)\s*"
        r"[-–—to]+\s*"
        r"(Present|Current|Now|\b(?:\d{1,2}/\d{4}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4})\b)\s*$",
        line,
        re.I,
    )
    if role_dates:
        return ("Company", role_dates.group(1).strip(), role_dates.group(2).strip(), role_dates.group(3).strip())

    dates = DATE_RANGE.search(line)
    start, end = "N/A", "Present"
    remainder = line

    if dates:
        start = dates.group(1).strip()
        end = dates.group(2).strip()
        remainder = line[: dates.start()] + line[dates.end() :]

    remainder = remainder.replace("|", " — ").strip(" -—–|,")

    if " — " in remainder:
        parts = [p.strip() for p in remainder.split(" — ") if p.strip()]
        company = parts[0]
        title = parts[1] if len(parts) > 1 else "Role"
    elif " at " in remainder.lower():
        idx = remainder.lower().index(" at ")
        title = remainder[:idx].strip()
        company = remainder[idx + 4 :].strip()
    elif "," in remainder:
        parts = [p.strip() for p in remainder.split(",", 1)]
        title = parts[0]
        company = parts[1] if len(parts) > 1 else "Company"
    else:
        company = remainder or "Company"
        title = "Role"

    return company, title, start, end


def _split_job_title_from_body(line: str) -> tuple[str, str]:
    """When PDF merges 'Company Title Designed...' into one line, split title from bullets."""
    match = re.match(
        r"^(.+?\b(?:Engineer|Developer|Manager|Analyst|Intern|Lead|Architect|Consultant|"
        r"Specialist|Associate|Director|Coordinator|Designer|Administrator|Scientist))\s+"
        r"(?=(?:Designed|Developed|Built|Created|Implemented|Managed|Led|Delivered|Configured|Applied))",
        line,
        re.I,
    )
    if match:
        return match.group(1).strip(), line[match.end() :].strip()
    return "", line


def _parse_company_title(combined: str) -> tuple[str, str]:
    """Best-effort split of 'Kore AI Software Engineer' -> company + title."""
    role_words = (
        r"Software Engineer|AI Engineer|Product Manager|Data Scientist|DevOps Engineer|"
        r"Full Stack Developer|Backend Developer|Frontend Developer|ML Engineer|Consultant"
    )
    role_match = re.search(role_words, combined, re.I)
    if role_match:
        company = combined[: role_match.start()].strip(" ,-|")
        title = role_match.group(0).strip()
        return company or combined, title
    parts = combined.rsplit(" ", 2)
    if len(parts) >= 2:
        return " ".join(parts[:-1]).strip(), parts[-1].strip()
    return combined, "Role"


def _extract_header(lines: list[str]) -> tuple[PersonalDetails, int]:
    personal = PersonalDetails()
    if not lines:
        return personal, 0

    name_line = lines[0]
    clean_name, _contact = split_name_contact_line(name_line)
    personal.name = clean_name or name_line
    personal = _parse_contact(name_line, personal)
    idx = 1

    while idx < min(len(lines), 6):
        line = lines[idx]
        if _match_section(line):
            break
        if EMAIL.search(line) or PHONE.search(line) or LINKEDIN.search(line):
            personal = _parse_contact(line, personal)
            idx += 1
            continue
        if not personal.title and len(line) < 80 and not _is_bullet(line):
            personal.title = line
            idx += 1
            continue
        break

    return personal, idx


def parse_resume_text(text: str, *, allow_empty: bool = False) -> ResumeData:
    """Parse resume text into structured JSON, preserving as much content as possible."""
    if not text or not text.strip():
        if allow_empty:
            return ResumeData()
        from app.services.exceptions import ResumeValidationError
        raise ResumeValidationError()

    lines = _expand_lines(text)
    personal, start_idx = _extract_header(lines)
    personal = _parse_contact(text, personal)

    summary_parts: list[str] = []
    experience: list[ExperienceItem] = []
    projects: list[ProjectItem] = []
    skills: list[str] = []
    education: list[EducationItem] = []
    certifications: list[str] = []
    awards: list[str] = []

    section: Optional[str] = None
    current_exp: Optional[ExperienceItem] = None
    current_project: Optional[ProjectItem] = None
    summary_lines: list[str] = []

    def flush_experience():
        nonlocal current_exp
        if current_exp and (current_exp.bullets or current_exp.company):
            experience.append(current_exp)
        current_exp = None

    def flush_project():
        nonlocal current_project
        if current_project and current_project.title:
            projects.append(current_project)
        current_project = None

    skip_next = 0
    for rel_idx, line in enumerate(lines[start_idx:]):
        if skip_next:
            skip_next -= 1
            continue

        header = _match_section(line)
        if header:
            flush_experience()
            flush_project()
            if header == "languages" and section == "education":
                continue
            section = header
            if section == "summary":
                summary_lines = []
            continue

        if section == "summary":
            summary_lines.append(line)
            continue

        if section == "skills":
            if match_section_line(line):
                continue
            if _looks_like_job_header(line) or _split_job_title_from_body(line)[0]:
                flush_experience()
                section = "experience"
            else:
                parts = re.split(r"[,;|•\n]|(?<=\))\s+(?=[A-Z])|(?:\s{2,})", line)
                skills.extend(s.strip() for s in parts if s.strip() and len(s.strip()) > 1)
                continue

        if section == "experience":
            if _is_bullet(line):
                if not current_exp:
                    current_exp = ExperienceItem(
                        company="Experience", title="Role", start_date="", end_date=""
                    )
                current_exp.bullets.append(_strip_bullet(line))
            elif _looks_like_job_header(line):
                flush_experience()
                company, title, start, end = _parse_job_header(line)
                if company == "Company" and title != "Role":
                    pending_company_title = (title, start, end)
                    next_idx = start_idx + rel_idx + 1
                    next_line = lines[next_idx] if next_idx < len(lines) else None
                    if next_line and not _is_bullet(next_line) and not _match_section(next_line):
                        company = next_line.strip()
                        skip_next = 1
                        title = pending_company_title[0]
                        start, end = pending_company_title[1], pending_company_title[2]
                    else:
                        company = title
                        title = "Role"
                current_exp = ExperienceItem(
                    company=company, title=title, start_date=start, end_date=end, bullets=[]
                )
            elif not current_exp:
                title_prefix, body = _split_job_title_from_body(line)
                if title_prefix and body:
                    company, title = _parse_company_title(title_prefix)
                    current_exp = ExperienceItem(
                        company=company,
                        title=title,
                        start_date="",
                        end_date="",
                        bullets=[],
                    )
                    for part in split_merged_accomplishments(body):
                        current_exp.bullets.append(part)
                elif not _is_bullet(line) and len(line) < 100:
                    next_idx = start_idx + rel_idx + 1
                    next_line = lines[next_idx] if next_idx < len(lines) else None
                    if next_line and _looks_like_job_header(next_line):
                        company, _, start, end = _parse_job_header(next_line)
                        current_exp = ExperienceItem(
                            company=company,
                            title=line,
                            start_date=start,
                            end_date=end,
                            bullets=[],
                        )
                        skip_next = 1
                    elif len(line) > 12:
                        current_exp = ExperienceItem(
                            company="Experience", title="Role", start_date="", end_date=""
                        )
                        for part in split_merged_accomplishments(line):
                            current_exp.bullets.append(part)
                    else:
                        flush_experience()
                        company, title, start, end = _parse_job_header(line)
                        current_exp = ExperienceItem(
                            company=company, title=title, start_date=start, end_date=end, bullets=[]
                        )
                elif len(line) > 12:
                    current_exp = ExperienceItem(
                        company="Experience", title="Role", start_date="", end_date=""
                    )
                    for part in split_merged_accomplishments(line):
                        current_exp.bullets.append(part)
            elif current_exp:
                title_is_placeholder = (
                    not current_exp.title.strip()
                    or current_exp.title in ("Role", "Company")
                    or current_exp.title == current_exp.company
                )
                if not current_exp.bullets and len(line) < 90 and title_is_placeholder:
                    current_exp.title = line
                else:
                    parts = (
                        [_strip_bullet(line)]
                        if _is_bullet(line)
                        else split_merged_accomplishments(line)
                    )
                    current_exp.bullets.extend(parts)
            continue

        if section == "languages":
            continue

        if section == "interests":
            awards.append(_strip_bullet(line))
            continue

        if section == "awards":
            if match_section_line(line):
                continue
            for part in split_merged_accomplishments(_strip_bullet(line)):
                if part and not match_section_line(part):
                    awards.append(part)
            continue

        if section == "projects":
            if current_project and _is_likely_project_title(line) and current_project.description:
                flush_project()
                current_project = ProjectItem(title=line, description="")
            elif current_project is None:
                current_project = ProjectItem(title=line, description="")
            elif line.lower().startswith("tech:"):
                parts = [t.strip() for t in line[5:].split(",") if t.strip()]
                current_project.tech_stack = parts
            else:
                part = _strip_bullet(line)
                bullets = normalize_bullet_list([part])
                for bullet in bullets:
                    current_project.description += (
                        ("\n" if current_project.description else "") + f"• {bullet}"
                    )
                url_match = URL_PATTERN.search(part)
                if url_match and not current_project.url:
                    current_project.url = url_match.group(0).rstrip(".,;")
            continue

        if section == "education":
            if line.upper().startswith("LANGUAGES"):
                continue
            if match_section_line(line) in ("education", "certifications", "languages"):
                continue
            education.append(_parse_education_line(line))
            continue

        if section == "certifications":
            certifications.append(_strip_bullet(line))
            continue

        # Unsectioned content before first header → summary; bullets → last job
        if _is_bullet(line) and experience:
            experience[-1].bullets.append(_strip_bullet(line))
        elif not section:
            summary_lines.append(line)

    flush_experience()
    flush_project()

    for exp in experience:
        exp.bullets = normalize_bullet_list(exp.bullets)

    summary = "\n".join(summary_lines).strip()

    if not experience:
        chunks = [
            l for l in lines[start_idx:]
            if len(l) > 15
            and not _match_section(l)
            and not EMAIL.search(l)
            and not PHONE.search(l)
            and l != personal.name
        ]
        if chunks:
            experience.append(
                ExperienceItem(
                    company="Experience",
                    title=personal.title or "Role",
                    start_date="",
                    end_date="",
                    bullets=chunks[:20],
                )
            )

    if not experience and not summary:
        body_lines = [
            l
            for l in lines[start_idx:]
            if l.strip()
            and len(l.strip()) > 3
            and not _match_section(l)
            and not EMAIL.search(l)
            and l != personal.name
        ]
        if body_lines:
            joined = " ".join(body_lines)[:1200].strip()
            if joined:
                summary = joined
                experience.append(
                    ExperienceItem(
                        company="Experience",
                        title=personal.title or "Role",
                        start_date="",
                        end_date="",
                        bullets=body_lines[:15],
                    )
                )

    return normalize_resume_data(
        ResumeData(
            personal=personal,
            summary=summary,
            experience=experience,
            projects=projects,
            education=education,
            skills=list(dict.fromkeys(skills)),
            certifications=certifications,
            awards=awards,
        )
    )


def _expand_lines(text: str) -> list[str]:
    """Preprocess wrapped PDF lines and split any inline multi-bullet text."""
    merged = preprocess_resume_text(text)
    expanded: list[str] = []
    for line in merged:
        if is_bullet_line(line):
            expanded.append(line)
            continue
        parts = split_inline_bullets(line)
        if len(parts) > 1:
            for part in parts:
                expanded.append(f"• {part}")
        else:
            expanded.append(line)
    return expanded


def _demo_resume() -> ResumeData:
    return ResumeData(
        personal=PersonalDetails(
            name="Alex Rivera",
            title="Product Operations Specialist",
            email="alex.rivera@email.com",
            location="San Francisco, CA",
            linkedin="linkedin.com/in/alexrivera",
        ),
        summary=(
            "Product Operations Specialist with 4+ years optimizing cross-functional workflows "
            "in fast-paced fintech environments. Skilled in process documentation, stakeholder "
            "communication, and resource planning."
        ),
        experience=[
            ExperienceItem(
                company="FinTech Solutions",
                title="Product Operations Specialist",
                start_date="2021",
                end_date="Present",
                bullets=[
                    "Managed communication between product, engineering, and design teams.",
                    "Wrote process documentation for onboarding new team members.",
                    "Assisted in tracking engineering bandwidth and sprint capacity.",
                ],
            )
        ],
        skills=["Project Management", "Documentation", "Cross-functional Communication", "Jira", "Confluence"],
        education=[
            EducationItem(institution="UC Berkeley", degree="B.A. Business Administration", year="2020")
        ],
    )
